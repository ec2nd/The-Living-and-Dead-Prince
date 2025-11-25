import re
import sys
import difflib
import random
import math
from datetime import datetime
from io import BytesIO
from pathlib import Path

# --- EXTERNAL LIBRARIES ---
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, ns
    from docx.oxml.ns import qn
    
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    
    import nltk
    from nltk.sentiment import SentimentIntensityAnalyzer
    
    from docx2pdf import convert as convert_to_pdf
except ImportError as e:
    print(f"❌ CRITICAL: Missing libraries. Run: pip install python-docx nltk matplotlib docx2pdf")
    sys.exit(1)

# --- CONFIGURATION ---
FILE_1 = "pg1232.txt"
FILE_2 = "pg57037.txt"
MASTER_DOCX_FILE = "master_running_report.docx"
FINAL_PDF_FILE = "The_Living_and_Dead_Prince.pdf"
ERROR_LOG_FILE = "compilation_errors.txt"
MIN_CONVERGENCE_WORDS = 2 
CHAPTERS_PER_REPORT = 5 

# --- NOVEL TITLE ---
NOVEL_TITLE_LINES = [
    "The Living",
    "and Dead",
    "Prince",
    "",
    "a collaborative novel",
    "",
    "by Nicolo Machiavelli,",
    "Luigi Ricci,",
    "W. K. Marriott,",
    "Gemini,",
    "and Edward Wells",
    " ",
    " ",
    "for nanogenmo 2025",
    " ",
    " ",
    "published by pleasepress"
]

# Regex to identify chapter headings
CHAPTER_PATTERN_STRING = r'^\s*(CHAPTER\s+(?:[IVXLDCM]+|\d+|.+?)).*$' 
CHAPTER_PATTERN = re.compile(CHAPTER_PATTERN_STRING, re.IGNORECASE | re.MULTILINE)

# --- NLTK SETUP ---
def ensure_nltk():
    resources = [('tokenizers/punkt', 'punkt'), 
                 ('taggers/averaged_perceptron_tagger', 'averaged_perceptron_tagger'), 
                 ('sentiment/vader_lexicon', 'vader_lexicon')]
    for path, name in resources:
        try: nltk.data.find(path)
        except LookupError: nltk.download(name, quiet=True)

ensure_nltk()
SIA = SentimentIntensityAnalyzer()

# --- NLP ANALYTICS ENGINE ---

def get_metrics(text):
    """
    Micro-level analytics for specific text blocks (Divergence Analysis).
    """
    pos_counts = {'Noun': 0, 'Verb': 0, 'Adj': 0}
    if not text.strip():
        return {'wc': 0, 'sent': 0, 'density': 0, 'avg_len': 0, 'pos': pos_counts}
    
    # Analyze tokens
    tokens = nltk.word_tokenize(text)
    words = [w.lower() for w in tokens if w.isalnum()]
    wc = len(words) # Total Word Count used for Magnitude
    
    # Sentiment of the WHOLE block
    sent = SIA.polarity_scores(text)['compound']
    
    unique = len(set(words))
    density = unique / wc if wc > 0 else 0
    avg_len = sum(len(w) for w in words) / wc if wc > 0 else 0
    
    tagged = nltk.pos_tag(tokens)
    for _, tag in tagged:
        if tag.startswith('N'): pos_counts['Noun'] += 1
        elif tag.startswith('V'): pos_counts['Verb'] += 1
        elif tag.startswith('J'): pos_counts['Adj'] += 1
        
    return {'wc': wc, 'sent': sent, 'density': density, 'avg_len': avg_len, 'pos': pos_counts}

# --- VISUALIZATION ENGINE ---

def calculate_bar_width(wc):
    """
    Calculates bar width based on Total Word Count (Magnitude).
    Base unit (1 word) is set to approx width of the Neutral Dot (0.05).
    Increases linearly per word.
    """
    if wc <= 0: return 0.05
    
    base_width = 0.05  # Approximate visual width of the 'dot'
    growth_per_word = 0.012 # Gentle linear growth
    
    width = base_width + (wc * growth_per_word)
    
    # Cap width at 0.9 to prevent overlapping into the next column too much
    return min(0.9, width)

def create_sentiment_chart(m1, m2):
    """
    Visualizes sentiment.
    - Bar Height: Sentiment Score
    - Bar Width: Magnitude (Total Word Count)
    - Zero Handling: Places a dot AND a horizontal line to show 'mass' even if score is 0.
    """
    fig, ax = plt.subplots(figsize=(3, 2))
    
    x_coords = [0, 1]
    labels = ['T1', 'T2']
    scores = [m1['sent'], m2['sent']]
    wcs = [m1['wc'], m2['wc']]
    
    widths = [calculate_bar_width(w) for w in wcs]
    colors = ['#4CAF50' if s >= 0 else '#F44336' for s in scores]
    
    # Plot Bars
    bars = ax.bar(x_coords, scores, color=colors, alpha=0.7, width=widths)
    
    ax.set_xticks(x_coords)
    ax.set_xticklabels(labels, fontsize=7)

    for i, score in enumerate(scores):
        # 1. Visualize Neutrality (0.0)
        # If score is effectively 0, the bar is invisible.
        # We plot a DOT to show "Neutral Sentiment"
        # We plot a HORIZONTAL LINE (hlines) to show "Word Volume"
        if abs(score) < 0.01:
            # The Dot
            ax.scatter(x_coords[i], 0, color='black', s=25, zorder=10, label='Neutral')
            
            # The Mass Line (Visualizes the width even though height is 0)
            half_w = widths[i] / 2
            ax.hlines(0, x_coords[i] - half_w, x_coords[i] + half_w, 
                      colors='black', linewidth=3, zorder=9, alpha=0.5)

    # 2. Add text labels for Word Count
    for i, rect in enumerate(bars):
        height = rect.get_height()
        width = rect.get_width()
        
        # Adjust label position
        y_pos = height + 0.15 if height >= 0 else height - 0.25
        if abs(height) < 0.1: y_pos = 0.25 
        
        ax.text(x_coords[i], y_pos,
                f"({wcs[i]}w)",
                ha='center', va='center', fontsize=6, color='black', alpha=0.8)

    ax.set_title('Sentiment + Vol (Width)', fontsize=8)
    ax.set_ylim(-1.3, 1.3)
    ax.axhline(0, color='black', linewidth=0.5)
    ax.tick_params(axis='y', which='major', labelsize=7)
    
    buf = BytesIO(); plt.tight_layout(); plt.savefig(buf, format='png', dpi=90); plt.close(fig); buf.seek(0)
    return buf

def create_density_chart(m1, m2):
    if m1['density'] == 0 and m2['density'] == 0: return None

    fig, ax = plt.subplots(figsize=(3, 2))
    labels = ['T1', 'T2']
    values = [m1['density'], m2['density']]
    
    ax.barh(labels, values, color=['#2196F3', '#FF9800'], alpha=0.7, height=0.5)
    
    ax.set_title('Lexical Density', fontsize=8)
    ax.set_xlim(0, 1)
    ax.tick_params(axis='both', which='major', labelsize=7)
    buf = BytesIO(); plt.tight_layout(); plt.savefig(buf, format='png', dpi=90); plt.close(fig); buf.seek(0)
    return buf

def create_pos_chart(m1, m2):
    v1, v2 = list(m1['pos'].values()), list(m2['pos'].values())
    if sum(v1) == 0 and sum(v2) == 0: return None
    
    fig, ax = plt.subplots(figsize=(3, 2))
    categories = ['Noun', 'Verb', 'Adj']
    v1 = [m1['pos'].get(k, 0) for k in categories]
    v2 = [m2['pos'].get(k, 0) for k in categories]
    x = range(len(categories)); width = 0.35
    ax.bar([i - width/2 for i in x], v1, width, label='T1', color='gray')
    ax.bar([i + width/2 for i in x], v2, width, label='T2', color='black')
    ax.set_xticks(x); ax.set_xticklabels(categories, fontsize=7)
    ax.set_title('Grammar', fontsize=8); ax.legend(fontsize=6)
    buf = BytesIO(); plt.tight_layout(); plt.savefig(buf, format='png', dpi=90); plt.close(fig); buf.seek(0)
    return buf

def generate_visual_analysis(doc, t1, t2):
    m1 = get_metrics(t1)
    m2 = get_metrics(t2)
    if m1['wc'] < 5 and m2['wc'] < 5: return

    delta_sent = abs(m1['sent'] - m2['sent']) / 2.0 
    delta_dens = abs(m1['density'] - m2['density']) 
    total_words = max(1, m1['wc'] + m2['wc'])
    pos_diff_sum = sum(abs(m1['pos'][k] - m2['pos'][k]) for k in m1['pos'])
    delta_pos = pos_diff_sum / total_words if total_words else 0
    
    metrics_map = {'Sentiment': delta_sent, 'Lexical Density': delta_dens, 'Grammar': delta_pos}
    chart_map = {'Sentiment': create_sentiment_chart, 'Lexical Density': create_density_chart, 'Grammar': create_pos_chart}

    div_feat = max(metrics_map, key=metrics_map.get)
    conv_feat = min(metrics_map, key=metrics_map.get)
    
    mode = random.choice(['div', 'conv', 'both'])
    charts_requested = []
    if mode == 'div': charts_requested.append(div_feat)
    elif mode == 'conv': charts_requested.append(conv_feat)
    else:
        charts_requested.append(div_feat)
        if conv_feat != div_feat: charts_requested.append(conv_feat)
            
    valid_images = []
    for feature_name in charts_requested:
        func = chart_map[feature_name]
        img_stream = func(m1, m2)
        if img_stream:
            valid_images.append(img_stream)
    
    if valid_images:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run("NLP Analysis: ")
        run.font.bold = True; run.font.size = Pt(9)
        
        narrative = ""
        # Check for Double Neutrality
        if abs(m1['sent']) < 0.01 and abs(m2['sent']) < 0.01:
            narrative += "Both translations exhibit neutral sentiment (0.0). "
        else:
            narrative += f"The feature of greatest divergence is {div_feat}. "
            
        narrative += f"The feature of greatest convergence is {conv_feat}."
        
        p.add_run(narrative).font.size = Pt(9)
        
        table = doc.add_table(rows=1, cols=len(valid_images))
        table.autofit = True
        remove_table_borders(table)
        for i, img_stream in enumerate(valid_images):
            cell = table.rows[0].cells[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].add_run().add_picture(img_stream, width=Inches(2.0))
        doc.add_paragraph()

# --- UTILITIES ---

def finalize_critical_error(message, details=""):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"\n❌ CRITICAL ERROR: {message}\nDETAILS: {details}")
    try: Path(ERROR_LOG_FILE).write_text(f"{timestamp}\n{message}\n{details}", encoding='utf-8')
    except: pass

def split_by_chapter_boundary(raw_text):
    matches = list(CHAPTER_PATTERN.finditer(raw_text))
    chapters = []
    preface = ""
    if not matches: return raw_text, []
    preface = raw_text[:matches[0].start()]
    for i in range(len(matches)):
        match = matches[i]
        title = match.group(1).strip()
        start = match.end()
        end = matches[i+1].start() if i + 1 < len(matches) else len(raw_text)
        content = raw_text[start:end].strip()
        chapters.append({'title': title, 'content': content})
    return preface, chapters

def get_tokens_with_indices(text):
    tokens = []
    for match in re.finditer(r"\b\w+(?:['’]\w+)*\b", text):
        tokens.append((match.group(0).lower(), match.start(), match.end()))
    return tokens

def find_lcs_match(s1, s2):
    t1 = get_tokens_with_indices(s1)
    t2 = get_tokens_with_indices(s2)
    if not t1 or not t2: return None
    w1 = [t[0] for t in t1]; w2 = [t[0] for t in t2]
    matcher = difflib.SequenceMatcher(None, w1, w2, autojunk=False)
    match = matcher.find_longest_match(0, len(w1), 0, len(w2))
    if match.size >= MIN_CONVERGENCE_WORDS:
        f1, l1 = t1[match.a], t1[match.a + match.size - 1]
        f2, l2 = t2[match.b], t2[match.b + match.size - 1]
        return (f1[1], l1[2], f2[1], l2[2])
    return None

# --- DOCX GENERATORS ---

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    if tblPr.find(qn('w:tblBorders')): tblPr.remove(tblPr.find(qn('w:tblBorders')))
    borders = OxmlElement('w:tblBorders')
    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{border}'); el.set(qn('w:val'), 'none'); borders.append(el)
    tblPr.append(borders)

def add_page_number_xml(run):
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def apply_mla_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'; style.font.size = Pt(12); style.font.color.rgb = RGBColor(0,0,0)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0); style.paragraph_format.space_before = Pt(0)
    try:
        h1 = doc.styles['Heading 1']
        h1.font.name = 'Times New Roman'; h1.font.size = Pt(14); h1.font.color.rgb = RGBColor(0,0,0)
        h1.font.bold = True; h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except: pass
    for section in doc.sections:
        p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number_xml(p.add_run())

# --- COVER ART ---
def generate_cover_image(sentiments):
    fig, ax = plt.subplots(figsize=(8.5, 11))
    x = range(len(sentiments)); y = sentiments
    for i in range(len(x) - 1):
        width = max(0.5, 26 - i) 
        ax.plot(x[i:i+2], y[i:i+2], color='black', linewidth=width, solid_capstyle='round', alpha=0.8)
    ax.scatter(x, y, color='#D32F2F', s=60, zorder=5, edgecolors='black')
    ax.text(0.92, 0.5, "the living and dead prince", transform=ax.transAxes,
            rotation=270, va='center', ha='center', fontsize=48, fontname='Times New Roman', color='#333333')
    authors = ["Nicolo", "Machiavelli", "Luigi", "Ricci", "W. K.", "Marriott", "Gemini", "", "Edward", "Wells"]
    y_pos = 0.9; step = 0.8 / len(authors)
    for name in authors:
        if name: ax.text(0.08, y_pos, name, transform=ax.transAxes, fontsize=36, alpha=0.15, ha='left', va='center', weight='bold', fontname='Arial')
        y_pos -= step
    ax.set_xlim(-1, len(sentiments)); ax.set_ylim(-1.1, 1.1); ax.axis('off')
    buf = BytesIO(); plt.savefig(buf, format='png', dpi=300); plt.close(fig); buf.seek(0)
    return buf

def force_toc_update(doc):
    """Injects the updateFields element into document settings to force TOC update on open."""
    settings = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)

def generate_toc(doc):
    # Field code logic from the user's provided snippet
    doc.add_paragraph("Table of Contents", style='Heading 1')
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    fldChar1.set(qn('w:dirty'), 'true') # Mark as dirty to force regen
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2); run._r.append(fldChar3)
    doc.add_page_break()

def generate_title_page(doc):
    for _ in range(4): doc.add_paragraph()
    for line in NOVEL_TITLE_LINES:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "Prince" in line or "Living" in line:
            p.runs[0].font.size = Pt(28); p.runs[0].font.bold = True
    doc.add_page_break()

def generate_editor_letter(doc):
    doc.add_paragraph(datetime.now().strftime("%d %B %Y")).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph('Herein is a completed draft manuscript of a touch analysis of two translations of')
    p.add_run(' The Prince ').font.italic = True
    p.add_run('by Nicolo Machiavelli, one by W. K. Marriott and the other by Luigi Ricci.')
    p = doc.add_paragraph('The work uses a chapter-bound convergence/divergence analysis. Convergent text blocks (>= 2 words) appear centered. Divergent text from the two translations of')
    p.add_run(' The Prince ').font.italic = True
    p.add_run('appears in two columns and are followed by NLP-based analysis.')
    doc.add_page_break()

def generate_source_title_content(doc, content):
    for line in content.split('\n'):
        p = doc.add_paragraph(line, style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs: run.font.size = Pt(16)
    doc.add_page_break()

# --- PROCESSING ---

def generate_convergence_block(doc, text, style='Normal'):
    if not text.strip(): return
    for b in re.split(r'\n\s*\n', text):
        if b.strip():
            p = doc.add_paragraph(b.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style = style

def generate_divergence_block(doc, t1, t2):
    if not re.search(r'\w', t1) and not re.search(r'\w', t2): return
    table = doc.add_table(rows=1, cols=2)
    remove_table_borders(table)
    table.autofit = False
    w = Inches(3.2)
    for cell in table.rows[0].cells: cell.width = w
    
    c1 = table.rows[0].cells[0]; c1.paragraphs[0].clear()
    for b in re.split(r'\n\s*\n', t1):
        if b.strip(): c1.add_paragraph(b.strip())
        
    c2 = table.rows[0].cells[1]; c2.paragraphs[0].clear()
    for b in re.split(r'\n\s*\n', t2):
        p = c2.add_paragraph(b.strip()); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
    generate_visual_analysis(doc, t1, t2)

def recursive_compare(doc, s1, s2):
    if not s1.strip() and not s2.strip(): return
    match = find_lcs_match(s1, s2)
    if match:
        start1, end1, start2, end2 = match
        recursive_compare(doc, s1[:start1], s2[:start2])
        generate_convergence_block(doc, s1[start1:end1])
        recursive_compare(doc, s1[end1:], s2[end2:])
    else:
        generate_divergence_block(doc, s1, s2)

def clean_residual_titles(text):
    lines = text.splitlines()
    while lines and (CHAPTER_PATTERN.match(lines[0].strip()) or not lines[0].strip()): lines.pop(0)
    return "\n".join(lines).strip()

def extract_subtitle(text):
    lines = text.splitlines()
    start = 0
    while start < len(lines) and not lines[start].strip(): start += 1
    if start < len(lines):
        first = lines[start].strip()
        if first.isupper() and len(first) > 5 and re.search(r'[A-Z]', first):
            sub = [first]; curr = start + 1
            while curr < len(lines):
                nxt = lines[curr].strip()
                if nxt.isupper() and re.search(r'[A-Z]', nxt): sub.append(nxt); curr += 1
                else: break
            return "\n".join(sub), "\n".join(lines[curr:]).strip()
    return None, text.strip()

def process_chapter(doc, title, c1, c2):
    generate_convergence_block(doc, title, style='Heading 1')
    doc.add_paragraph()
    
    c1 = clean_residual_titles(c1); c2 = clean_residual_titles(c2)
    s1, b1 = extract_subtitle(c1); s2, b2 = extract_subtitle(c2)
    
    if s1 or s2:
        if s1 and s2 and s1 == s2: generate_convergence_block(doc, s1)
        else: recursive_compare(doc, s1 or "", s2 or "")
        doc.add_paragraph()
        
    try: recursive_compare(doc, b1, b2)
    except RecursionError: generate_divergence_block(doc, b1, b2)

def save_final(doc):
    try:
        force_toc_update(doc)
        doc.save(MASTER_DOCX_FILE)
        print(f"✅ DOCX Saved: {MASTER_DOCX_FILE}")
        print(f"⏳ Converting to PDF...")
        convert_to_pdf(MASTER_DOCX_FILE, FINAL_PDF_FILE)
        print(f"✅ PDF Saved: {FINAL_PDF_FILE}")
    except Exception as e: finalize_critical_error("Save Failed", str(e))

CHAPTER_SENTIMENTS = []

def process_texts():
    print("--- Starting Visual Analytics Analysis + PDF ---")
    try: raw1 = Path(FILE_1).read_text(encoding='utf-8'); raw2 = Path(FILE_2).read_text(encoding='utf-8')
    except Exception as e: finalize_critical_error("Read Error", str(e)); return

    _, ch1 = split_by_chapter_boundary(raw1)
    _, ch2 = split_by_chapter_boundary(raw2)
    
    # 1. Macro Sentiment Analysis (Pre-Pass) for Cover Art
    print("📊 Analyzing Sentiment Trajectory for Cover...")
    for i in range(max(len(ch1), len(ch2))):
        t1 = ch1[i]['content'] if i < len(ch1) else ""
        t2 = ch2[i]['content'] if i < len(ch2) else ""
        CHAPTER_SENTIMENTS.append(SIA.polarity_scores(t1 + " " + t2)['compound'])
    
    cover_img = generate_cover_image(CHAPTER_SENTIMENTS)
    
    doc = Document()
    apply_mla_style(doc)
    
    # Insert Cover
    doc.add_picture(cover_img, width=Inches(6.5))
    doc.add_page_break()
    
    generate_title_page(doc)
    generate_toc(doc)
    
    p1, _ = split_by_chapter_boundary(raw1)
    if p1.strip(): generate_source_title_content(doc, p1); doc.add_page_break()
    generate_editor_letter(doc)

    max_len = max(len(ch1), len(ch2))

    for i in range(max_len):
        c1 = ch1[i] if i < len(ch1) else {'title':'', 'content':''}
        c2 = ch2[i] if i < len(ch2) else {'title':'', 'content':''}
        title = c1['title'] if c1['title'] else c2['title']
        if not title: title = f"CHAPTER {i+1}"
        
        print(f"Processing {title.strip()}...")
        if i > 0: doc.add_page_break()
        process_chapter(doc, title, c1['content'], c2['content'])

    save_final(doc)
    print("\nProcessing Complete.")

if __name__ == "__main__":
    process_texts()
